from docx import Document
import os

# print(os.getcwd())
doc = Document('22215150135_何超华_任务书.docx')

# 读取段落
for p in doc.paragraphs:
    print(p.text)

# 读取表格
# 一个表格
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            print(c.text)
