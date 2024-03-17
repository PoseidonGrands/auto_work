from docx import Document
from docx.shared import Inches

doc = Document()

# 添加标题
title = doc.add_heading('我是标题', 0)
title.add_run('哈哈哈')

# 添加段落
p = doc.add_paragraph('段落1')
p.add_run('...')

# 添加图片
image = doc.add_picture('1.jpeg', width=Inches(3))

# 添加表格
table = doc.add_table(rows=6, cols=6)

titles = ['name', 'age', 'sex']
data = [
    ['xiaohong', '10', 'boy'],
    ['xiaohei', '11', 'girl'],
    ['john', '16', 'boy']
]

title_cells = table.rows[0].cells
title_cells[0].text = titles[0]
title_cells[1].text = titles[1]
title_cells[2].text = titles[2]

for r_index in range(1, 4):
    _cells = table.rows[r_index].cells
    for c_index in range(3):
        _cells[c_index].text = data[r_index - 1][c_index]

# 添加新页
doc.add_page_break()

doc.save('test.docx')
