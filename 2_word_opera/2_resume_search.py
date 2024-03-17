import glob
from docx import Document


class ReadDoc(object):
    def __init__(self, path):
        self.doc = Document(path)

        self.p_text = ''
        self.table_text = ''

        self.get_para()
        self.get_table()

    def get_para(self):
        for p in self.doc.paragraphs:
            self.p_text += p.text + '\n '

    def get_table(self):
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.table_text += cell.text + ','
                self.table_text += '\n'


def search_word(path, targets):
    files = glob.glob(path)
    res = []

    for file in files:
        if glob.os.path.isfile(file):
            if file.endswith('docx'):
                doc = ReadDoc(file)
                contents = doc.p_text + doc.table_text

                is_target = True
                for word in targets:
                    if word not in contents:
                        is_target = False
                        break
                if is_target:
                    res.append(file)
    return res
 

if __name__ == '__main__':
    path = glob.os.path.join(glob.os.getcwd(), '*')

    res = search_word(path, ['python'])
    print(res)
