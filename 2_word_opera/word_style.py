import random

from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# 1、添加全局样式
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.color.rgb = RGBColor(255, 0, 0)
style.font.size = Pt(16)

# 添加标题
title = doc.add_heading('', 0)
_t = title.add_run('我是标题哈哈哈')
# 2、添加文本样式
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title.style.font.size = Pt(12)
# 某些只可以对追加的内容添加（原标题可以留空，然后全部内容都追加即可解决原标题无法设置
_t.italic = True

# 添加段落
p = doc.add_paragraph('段落1')
p.add_run('...')

# 3、图片样式
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
_p = p.add_run()
image = _p.add_picture('1.jpeg', width=Inches(3))

# 4、表格样式
# 获取全部表格样式
table_styles = []
for style in doc.styles:
    if style.type == WD_STYLE_TYPE.TABLE:
        table_styles.append(style.name)
random_style = random.choice(table_styles)
table = doc.add_table(rows=3, cols=3, style=random_style)

doc.save('test_2.docx')
